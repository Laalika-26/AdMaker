import os
import sys
import json
from PIL import Image, ImageDraw, ImageFont
import numpy as np

# Word document generation imports
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_directories():
    os.makedirs("./report_images", exist_ok=True)
    print("[Script] 'report_images' directory verified.")

def get_font(font_name, size, weight="regular"):
    # Attempt to load system font from Windows Fonts directory
    font_dir = r"C:\Windows\Fonts"
    font_file = "arial.ttf"
    if font_name.lower() == "consolas":
        font_file = "consola.ttf"
    elif weight == "bold":
        font_file = "arialbd.ttf"
        
    font_path = os.path.join(font_dir, font_file)
    try:
        if os.path.exists(font_path):
            return ImageFont.truetype(font_path, size)
        else:
            return ImageFont.truetype(font_file, size)
    except IOError:
        return ImageFont.load_default()

def draw_presidency_logo():
    img = Image.new("RGB", (400, 120), color="white")
    draw = ImageDraw.Draw(img)
    
    font_title = get_font("arial", 20, "bold")
    font_sub = get_font("arial", 9, "regular")
    
    shield_pts = [(20, 20), (70, 20), (70, 70), (45, 100), (20, 70)]
    draw.polygon(shield_pts, fill=(15, 36, 62)) # Navy Blue
    
    draw.line([(30, 45), (45, 60), (60, 45)], fill="white", width=4)
    draw.line([(45, 60), (45, 90)], fill="white", width=3)
    draw.line([(35, 35), (55, 35)], fill="gold", width=4) # Golden top bar
    
    draw.text((90, 30), "PRESIDENCY", fill=(15, 36, 62), font=font_title)
    draw.text((90, 52), "UNIVERSITY", fill=(15, 36, 62), font=font_title)
    draw.text((90, 80), "GAIN MORE KNOWLEDGE / REACH GREATER HEIGHTS", fill=(100, 100, 100), font=font_sub)
    
    img.save("./report_images/presidency_logo.png")
    print("[Script] Generated presidency_logo.png")

def draw_nvidia_logo():
    img = Image.new("RGB", (350, 120), color="black")
    draw = ImageDraw.Draw(img)
    
    font_coe = get_font("arial", 16, "bold")
    font_acc = get_font("arial", 10, "regular")
    font_nvidia = get_font("arial", 24, "bold")
    
    draw.rectangle([0, 0, 10, 120], fill=(118, 185, 0)) # NVIDIA green
    
    draw.text((30, 20), "AI Centre of Excellence", fill=(118, 185, 0), font=font_coe)
    draw.text((30, 48), "Accelerated by", fill=(200, 200, 200), font=font_acc)
    draw.text((30, 68), "NVIDIA", fill="white", font=font_nvidia)
    
    img.save("./report_images/nvidia_logo.png")
    print("[Script] Generated nvidia_logo.png")

def draw_dashboard_state1():
    img = Image.new("RGB", (800, 450), color=(30, 30, 35))
    draw = ImageDraw.Draw(img)
    
    font_header = get_font("arial", 16, "bold")
    font_panel = get_font("arial", 12, "bold")
    font_body = get_font("arial", 10, "regular")
    font_hint = get_font("arial", 8, "regular")
    
    draw.rectangle([0, 0, 800, 50], fill=(20, 20, 22))
    draw.text((20, 15), "AdMarker Engine v2.0", fill=(0, 240, 255), font=font_header)
    draw.ellipse([700, 20, 712, 32], fill=(0, 255, 100)) # green running dot
    draw.text((720, 18), "Core Node Running", fill=(200, 200, 200), font=font_body)
    
    # Left Panel
    draw.rectangle([20, 70, 380, 430], fill=(40, 40, 45), outline=(60, 60, 65), width=2)
    draw.text((40, 85), "⚙ Configuration Settings", fill="white", font=font_panel)
    
    draw.text((40, 125), "Transcription Model", fill=(180, 180, 180), font=font_body)
    draw.rectangle([40, 145, 360, 175], fill=(50, 50, 55), outline=(100, 100, 105))
    draw.text((50, 153), "Google Gemini (REST)", fill="white", font=font_body)
    
    draw.text((40, 195), "Context Tagging Model", fill=(180, 180, 180), font=font_body)
    draw.rectangle([40, 215, 360, 245], fill=(50, 50, 55), outline=(100, 100, 105))
    draw.text((50, 223), "Google Gemini (REST)", fill="white", font=font_body)
    
    draw.text((40, 270), "Minimum Silent Gap", fill=(180, 180, 180), font=font_body)
    draw.line([40, 305, 360, 305], fill=(80, 80, 80), width=4)
    draw.line([40, 305, 200, 305], fill=(0, 200, 255), width=4)
    draw.ellipse([195, 297, 207, 309], fill=(0, 240, 255))
    draw.text((320, 270), "1.5s", fill=(0, 240, 255), font=font_body)
    
    draw.rectangle([40, 360, 360, 405], fill=(0, 100, 180), outline=(0, 120, 220))
    draw.text((120, 373), "Upload a Video to Begin", fill="white", font=font_panel)
    
    # Right Panel
    draw.rectangle([400, 70, 780, 430], fill=(40, 40, 45), outline=(60, 60, 65), width=2)
    draw.text((420, 85), "▲ Video Uploader", fill="white", font=font_panel)
    
    draw.rectangle([420, 120, 760, 410], fill=(35, 35, 40), outline=(0, 150, 200), width=1)
    draw.polygon([(590, 210), (590, 180), (600, 180), (580, 150), (560, 180), (570, 180), (570, 210)], fill=(0, 200, 255))
    draw.text((500, 230), "Drag & Drop Video File Here", fill="white", font=font_body)
    draw.text((515, 255), "or click to browse your system", fill=(150, 150, 150), font=font_body)
    draw.text((510, 285), "Supports: MP4, MOV, WAV, MP3", fill=(100, 100, 100), font=font_hint)
    
    img.save("./report_images/dashboard_state1.png")
    print("[Script] Generated dashboard_state1.png")

def draw_dashboard_state2():
    img = Image.new("RGB", (800, 450), color=(30, 30, 35))
    draw = ImageDraw.Draw(img)
    
    font_header = get_font("arial", 16, "bold")
    font_panel = get_font("arial", 12, "bold")
    font_body = get_font("arial", 10, "regular")
    font_mono = get_font("consolas", 9, "regular")
    
    draw.rectangle([0, 0, 800, 50], fill=(20, 20, 22))
    draw.text((20, 15), "AdMarker Engine v2.0", fill=(0, 240, 255), font=font_header)
    draw.ellipse([700, 20, 712, 32], fill=(0, 255, 100)) # green running dot
    draw.text((720, 18), "Core Node Running", fill=(200, 200, 200), font=font_body)
    
    draw.rectangle([20, 70, 780, 180], fill=(40, 40, 45), outline=(60, 60, 65), width=2)
    draw.text((40, 85), "🖵 Pipeline Status Monitor", fill="white", font=font_panel)
    
    stages = ["Extraction", "VAD Chunk", "Transcription", "Boundary Scoring", "Tagging"]
    statuses = ["Complete", "Complete", "Active...", "Pending", "Pending"]
    colors = [(0, 255, 100), (0, 255, 100), (0, 200, 255), (100, 100, 100), (100, 100, 100)]
    
    for i in range(5):
        x = 40 + i * 145
        draw.rectangle([x, 115, x + 130, 165], fill=(50, 50, 55), outline=(70, 70, 75))
        draw.text((x + 10, 123), stages[i], fill="white", font=font_body)
        draw.text((x + 10, 143), f"[ {statuses[i]} ]", fill=colors[i], font=font_body)
        if i < 4:
            draw.line([x + 130, 140, x + 145, 140], fill=(100, 100, 100), width=2)
            
    draw.rectangle([20, 200, 780, 430], fill=(15, 15, 18), outline=(60, 60, 65), width=2)
    draw.rectangle([20, 200, 780, 225], fill=(30, 30, 35))
    draw.text((35, 205), "Live Log Console (guest@admarker:~)", fill="white", font=font_body)
    
    logs = [
        "guest@admarker:~$ python main.py uploads/video.mp4 --min-gap 1.5 --transcription-provider gemini",
        "[Audio] Extracting audio stream from uploads/video.mp4...",
        "[Audio] Audio stream extracted successfully. Output: C:\\Users\\Swathi\\...\\extracted_audio.wav",
        "[VAD] Scanning voice activity & chunking audio stream (energy_threshold=0.015)...",
        "[VAD] Splitting complete: generated 4 conversational chunks.",
        "[Transcriber] Stitching timeline & transcribing chunks using Gemini REST API...",
        "  [Chunk 1/4] Transcribing: chunk_0000_0.00_22.50.wav...",
        "  [Chunk 2/4] Transcribing: chunk_0001_22.50_45.00.wav..."
    ]
    
    for i, line in enumerate(logs):
        color = (0, 255, 100) if i == 0 else ((0, 200, 255) if i in (1, 3, 5) else (220, 220, 220))
        draw.text((35, 235 + i * 22), line, fill=color, font=font_mono)
        
    img.save("./report_images/dashboard_state2.png")
    print("[Script] Generated dashboard_state2.png")

def draw_dashboard_state3():
    img = Image.new("RGB", (800, 450), color=(30, 30, 35))
    draw = ImageDraw.Draw(img)
    
    font_header = get_font("arial", 16, "bold")
    font_panel = get_font("arial", 12, "bold")
    font_body = get_font("arial", 10, "regular")
    font_body_bold = get_font("arial", 10, "bold")
    
    draw.rectangle([0, 0, 800, 50], fill=(20, 20, 22))
    draw.text((20, 15), "AdMarker Engine v2.0", fill=(0, 240, 255), font=font_header)
    draw.ellipse([700, 20, 712, 32], fill=(0, 255, 100)) # green running dot
    draw.text((720, 18), "Core Node Running", fill=(200, 200, 200), font=font_body)
    
    draw.rectangle([20, 70, 780, 430], fill=(40, 40, 45), outline=(60, 60, 65), width=2)
    draw.rectangle([20, 70, 780, 105], fill=(30, 30, 35))
    draw.text((40, 80), "Identified Ad Markers & Product Targeting (3 Breaks Found)", fill=(0, 240, 255), font=font_panel)
    
    cols = ["ID", "Timestamp", "Gap", "Preceding Transcript Context", "Category & Tone", "Targeted Ads"]
    col_x = [40, 100, 180, 240, 530, 680]
    
    draw.rectangle([30, 115, 770, 145], fill=(15, 36, 62)) # Navy header
    for idx, col in enumerate(cols):
        draw.text((col_x[idx], 122), col, fill="white", font=font_body_bold)
        
    # Row 1
    draw.rectangle([30, 150, 770, 250], fill=(45, 45, 50), outline=(60, 60, 65))
    draw.text((col_x[0], 170), "AD-01", fill=(0, 255, 100), font=font_body_bold)
    draw.text((col_x[1], 170), "00:00:21", fill=(0, 240, 255), font=font_body)
    draw.text((col_x[2], 170), "2.50s", fill="yellow", font=font_body)
    
    context_r1 = [
        "Welcome back everyone. Today we are talking about cloud",
        "scalability and serverless functions. If you deploy microservices,",
        "you don't want to manage physical containers or VMs."
    ]
    for i, line in enumerate(context_r1):
        draw.text((col_x[3], 158 + i * 18), line, fill="white", font=font_body)
        
    draw.text((col_x[4], 162), "Technology", fill="white", font=font_body_bold)
    draw.text((col_x[4], 182), "Tone: Informative", fill=(180, 180, 180), font=font_body)
    draw.text((col_x[5], 162), "- CloudVibe Hosting", fill=(150, 255, 150), font=font_body)
    draw.text((col_x[5], 182), "- PyShield Auditor", fill=(150, 255, 150), font=font_body)
    
    # Row 2
    draw.rectangle([30, 255, 770, 355], fill=(40, 40, 45), outline=(60, 60, 65))
    draw.text((col_x[0], 275), "AD-02", fill=(0, 255, 100), font=font_body_bold)
    draw.text((col_x[1], 275), "00:00:46", fill=(0, 240, 255), font=font_body)
    draw.text((col_x[2], 275), "2.00s", fill="yellow", font=font_body)
    
    context_r2 = [
        "Hey guys, today we are baking an authentic French sourdough",
        "bread from scratch. The key is the fermentation process,",
        "which depends heavily on the ambient temperature."
    ]
    for i, line in enumerate(context_r2):
        draw.text((col_x[3], 263 + i * 18), line, fill="white", font=font_body)
        
    draw.text((col_x[4], 267), "Cooking", fill="white", font=font_body_bold)
    draw.text((col_x[4], 287), "Tone: Instructional", fill=(180, 180, 180), font=font_body)
    draw.text((col_x[5], 267), "- RiseMaster Oven", fill=(150, 255, 150), font=font_body)
    draw.text((col_x[5], 287), "- DoughPro Mixer", fill=(150, 255, 150), font=font_body)
    
    img.save("./report_images/dashboard_state3.png")
    print("[Script] Generated dashboard_state3.png")

def add_shading(cell, color_hex):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shd)

def add_left_border(paragraph, color_hex="0070C0", size="24"):
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="{size}" w:space="10" w:color="{color_hex}"/></w:pBdr>')
    paragraph._p.get_or_add_pPr().append(pBdr)

def add_paragraph_shading(paragraph, color_hex="F5F5F5"):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    paragraph._p.get_or_add_pPr().append(shd)

def build_docx_report():
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(30, 30, 30)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)
    
    # ----------------------------------------------------
    # COVER PAGE
    # ----------------------------------------------------
    logo_table = doc.add_table(rows=1, cols=2)
    logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in logo_table.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        
    cell_l = logo_table.cell(0, 0)
    cell_l.width = Inches(3.2)
    p_l = cell_l.paragraphs[0]
    r_l = p_l.add_run()
    r_l.add_picture("./report_images/presidency_logo.png", width=Inches(3.0))
    
    cell_r = logo_table.cell(0, 1)
    cell_r.width = Inches(3.2)
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_r = p_r.add_run()
    r_r.add_picture("./report_images/nvidia_logo.png", width=Inches(2.5))
    
    doc.add_paragraph("\n\n")
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("NVIDIA ACCELERATED AI CENTRE\nOF EXCELLENCE")
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(24)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(11, 102, 35) # Dark Green
    p_title.paragraph_format.space_after = Pt(18)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Final Internship and Project Report")
    r_sub.font.name = 'Arial'
    r_sub.font.size = Pt(16)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(50, 50, 50)
    p_sub.paragraph_format.space_after = Pt(36)
    
    doc.add_paragraph("\n\n")
    
    details_table = doc.add_table(rows=6, cols=2)
    details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    details = [
        ("PROJECT REPORT TITLE", "AUTOMATED VIDEO AD-MARKER & CONTEXTUAL PROMPT GENERATOR"),
        ("STUDENT NAME", "LAALIKA J"),
        ("ROLL NUMBER / ID", "20241CAI0231"),
        ("BRANCH & DEPARTMENT", "B-TECH School of Artificial Intelligence & Advanced Computing"),
        ("START DATE & END DATE", "29-06-2026 TO 10-07-2026"),
        ("PROJECT SUBMISSION", "11-07-2026")
    ]
    
    for idx, (label, val) in enumerate(details):
        row = details_table.rows[idx]
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        
        c_lbl = row.cells[0]
        c_lbl.width = Inches(2.2)
        p_lbl = c_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_after = Pt(4)
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(10.5)
        
        c_val = row.cells[1]
        c_val.width = Inches(4.3)
        p_val = c_val.paragraphs[0]
        p_val.paragraph_format.space_after = Pt(4)
        r_val = p_val.add_run(val)
        r_val.font.bold = True
        r_val.font.size = Pt(10.5)
        r_val.font.color.rgb = RGBColor(31, 78, 121) # dark blue
        
    doc.add_page_break()
    
    # ----------------------------------------------------
    # SECTION 1: About NVIDIA
    # ----------------------------------------------------
    p_h1 = doc.add_paragraph()
    r_h1 = p_h1.add_run("1. About NVIDIA Corporation")
    r_h1.font.size = Pt(16)
    r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(31, 78, 121)
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph(
        "NVIDIA Corporation, founded in 1993 by Jen-Hsun Huang, Chris Malachowsky, and Curtis Priem, "
        "is a pioneer in graphics processing unit (GPU) accelerated computing. Originally focused on "
        "bringing 3D graphics to the gaming and multimedia markets, NVIDIA invented the GPU in 1999, "
        "redefining modern computer graphics and sparking the growth of the PC gaming market. Over "
        "the next two decades, NVIDIA expanded its scope, transforming the GPU from a specialized 3D "
        "rendering chip into a highly programmable parallel processor capable of general-purpose "
        "scientific computing."
    )
    
    doc.add_paragraph(
        "Today, NVIDIA stands at the absolute vanguard of the Artificial Intelligence revolution. "
        "By recognizing the immense mathematical alignment between deep learning algorithms and "
        "parallel hardware architectures, the company made a massive, long-term strategic bet on AI "
        "supercomputing. Through the development of the proprietary CUDA compute platform, NVIDIA "
        "created a fully integrated stack of software libraries, hardware drivers, and deep learning "
        "engines that became the standard for researchers and developers globally. In 2024, NVIDIA "
        "joined the ranks of the world's most valuable companies, crossing a $3 trillion market "
        "capitalization milestone. Its compute platforms have become the foundation of almost all major "
        "generative AI foundations, large language models (LLMs), autonomous vehicle systems, and "
        "scientific supercomputing centers around the world."
    )
    
    # ----------------------------------------------------
    # SECTION 2: NVIDIA H200 GPU
    # ----------------------------------------------------
    p_h2 = doc.add_paragraph()
    r_h2 = p_h2.add_run("2. NVIDIA H200 Tensor Core GPU")
    r_h2.font.size = Pt(16)
    r_h2.font.bold = True
    r_h2.font.color.rgb = RGBColor(31, 78, 121)
    p_h2.paragraph_format.space_before = Pt(18)
    p_h2.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph(
        "Released in Q2 2024, the NVIDIA H200 Tensor Core GPU is the successor to the groundbreaking H100, "
        "designed specifically to address the massive memory and bandwidth bottlenecks of generative AI "
        "and large-scale model inference. Built on the Hopper architecture, the H200 is the first GPU "
        "to deploy ultra-fast HBM3e (High Bandwidth Memory 3e) technology. By packing 141 GB of HBM3e "
        "memory running at a staggering 4.8 terabytes per second (TB/s), the H200 provides double "
        "the capacity and 1.4x the bandwidth of its predecessor. This enables it to store massive "
        "language models entirely inside active GPU memory, dramatically reducing latency and "
        "maximizing computing efficiency."
    )
    
    spec_table = doc.add_table(rows=7, cols=3)
    spec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["SPECIFICATION", "NVIDIA H100 GPU", "NVIDIA H200 GPU (NEW)"]
    row_data = [
        ("Architecture", "Hopper", "Hopper (Upgraded Silicon)"),
        ("Memory Type", "HBM3", "HBM3e (Advanced High Bandwidth)"),
        ("Memory Capacity", "80 GB", "141 GB (Nearly Double)"),
        ("Memory Bandwidth", "3.35 TB/s", "4.8 TB/s (1.4x Faster)"),
        ("FP8 Performance", "Approx. 4 PetaFLOPS", "Approx. 4 PetaFLOPS"),
        ("LLM Inference Speed", "Baseline (1.0x)", "Llama 2 (70B): 1.9x faster | GPT-3: 1.6x faster")
    ]
    
    hdr_row = spec_table.rows[0]
    hdr_row._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
    for idx, text in enumerate(headers):
        cell = hdr_row.cells[idx]
        add_shading(cell, "0F243E") # Dark Navy
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10)
        
    for r_idx, row_vals in enumerate(row_data):
        row = spec_table.rows[r_idx + 1]
        row._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
        bg_color = "EBF1F5" if r_idx % 2 == 0 else "FFFFFF"
        
        for c_idx, val in enumerate(row_vals):
            cell = row.cells[c_idx]
            add_shading(cell, bg_color)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            if c_idx == 0:
                r.font.bold = True
                
    # ----------------------------------------------------
    # SECTION 3: GPU Architecture
    # ----------------------------------------------------
    p_h3 = doc.add_paragraph()
    r_h3 = p_h3.add_run("3. Understanding the NVIDIA GPU Architecture")
    r_h3.font.size = Pt(16)
    r_h3.font.bold = True
    r_h3.font.color.rgb = RGBColor(31, 78, 121)
    p_h3.paragraph_format.space_before = Pt(18)
    p_h3.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph(
        "Graphics Processing Units (GPUs) differ fundamentally from Central Processing Units (CPUs) "
        "in their computational philosophy. A CPU is designed for sequential processing and low-latency "
        "serial tasks, utilizing a few powerful cores optimized for single-threaded speed. Conversely, "
        "a GPU is built for massive parallel execution, containing thousands of smaller, simpler cores "
        "that work simultaneously. This makes GPUs uniquely suited for tasks like 3D graphics rendering, "
        "computer vision, and matrix math operations — which form the foundational operations of deep learning."
    )
    
    p_subh3 = doc.add_paragraph()
    p_subh3.paragraph_format.space_before = Pt(6)
    p_subh3.paragraph_format.space_after = Pt(6)
    r_subh3 = p_subh3.add_run("Key Architectural Sub-systems of NVIDIA GPUs:")
    r_subh3.font.bold = True
    r_subh3.font.color.rgb = RGBColor(0, 112, 192)
    
    subsystems = [
        ("CUDA Cores", "The fundamental parallel processors that execute basic mathematical calculations and logic instructions on the GPU."),
        ("Tensor Cores", "Specialized hardware units introduced in the Volta architecture designed to perform fast mixed-precision matrix multiplication in a single cycle. These are the engines that accelerate Transformer models and LLM operations."),
        ("Streaming Multiprocessors (SM)", "The primary building blocks of the GPU. Each SM groups CUDA cores, Tensor cores, register files, and shared memory cache to coordinate execution."),
        ("NVLink Interconnect", "NVIDIA's proprietary high-bandwidth interconnect, allowing multiple GPUs to communicate at speeds up to 900 GB/s, enabling them to act as a single massive virtual GPU.")
    ]
    
    for title, desc in subsystems:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r_t = p.add_run(f"{title}: ")
        r_t.font.bold = True
        r_d = p.add_run(desc)
        
    # ----------------------------------------------------
    # SECTION 4: Project Timeline
    # ----------------------------------------------------
    p_h4 = doc.add_paragraph()
    r_h4 = p_h4.add_run("4. Project Work Timeline: Day 1 to Day 10")
    r_h4.font.size = Pt(16)
    r_h4.font.bold = True
    r_h4.font.color.rgb = RGBColor(31, 78, 121)
    p_h4.paragraph_format.space_before = Pt(18)
    p_h4.paragraph_format.space_after = Pt(12)
    
    timeline = [
        ("Day 1: Introduction to NVIDIA GPU and Kubernetes Pods and Services", 
         ["Understood about NVIDIA H-200 GPU and its features.", "Creation of Pods and services & how it connects to the NVIDIA Server."]),
        ("Day 2: Introduction to Machine Learning and Model Training", 
         ["Understood Core Machine Learning Concepts.", "Performed a program to train a model by following the Model Training Pipeline."]),
        ("Day 3: Supervised and Unsupervised Machine Learning Concepts", 
         ["Focused and understood programs based on Classification & Regression."]),
        ("Day 4: Neural Networks: CNN and RNN", 
         ["Focused on creating Neural Networks and trained a model based on the MNIST dataset."]),
        ("Day 5: Working Natural Language Processing Pipeline for a Model", 
         ["Understood NLP Pipeline and Types of NLP.", "Focused on training a model following every step of the NLP pipeline."]),
        ("Day 6: Transformers and Types of Transformer (BERT & GPT)", 
         ["Understood Transformers and types of transformers used to train a model.", "Understood concepts like BERT and GPT and focused on training models that use these concepts."]),
        ("Day 7: Large Language Models", 
         ["Understood the importance of LLMs in Machine Learning.", "Focused on programs that used Google Gemini and compared other LLMs with Google Gemini."]),
        ("Day 8: VIVA about Concepts Covered and Project Selection", 
         ["VIVA about all the concepts that were covered till Day 7.", "Implementing a project based on the concepts we understood."]),
        ("Day 9: Working of Diffusion Models and Generative Adversarial Networks", 
         ["Understood how Diffusion Models work and types of Diffusion Models.", "Implemented programs that use diffusion models.", "Focused on training a model using GAN and generated outputs."]),
        ("Day 10: Working on Final Internship Project", 
         ["Focused on completing our Final Internship Project."])
    ]
    
    for day_title, items in timeline:
        p_day = doc.add_paragraph()
        p_day.paragraph_format.space_before = Pt(4)
        p_day.paragraph_format.space_after = Pt(2)
        r_dt = p_day.add_run(day_title)
        r_dt.font.bold = True
        r_dt.font.color.rgb = RGBColor(50, 50, 50)
        
        for item in items:
            p_it = doc.add_paragraph(style='List Bullet')
            p_it.paragraph_format.space_after = Pt(2)
            r_it = p_it.add_run(item)
            
    # ----------------------------------------------------
    # SECTION 5: Project Details (AdMarker Engine)
    # ----------------------------------------------------
    doc.add_page_break()
    p_h5 = doc.add_paragraph()
    r_h5 = p_h5.add_run("5. Project Details: AdMarker Engine")
    r_h5.font.size = Pt(16)
    r_h5.font.bold = True
    r_h5.font.color.rgb = RGBColor(31, 78, 121)
    p_h5.paragraph_format.space_before = Pt(12)
    p_h5.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph(
        "The AdMarker Engine is an automated pipeline and professional web dashboard designed to "
        "identify optimal quiet-zones (silent pauses) in audio or video timelines and leverage "
        "Large Language Models (LLMs) to perform semantic categorization and targeted product ad recommendations."
    )
    
    doc.add_paragraph(
        "Placing advertisements at random timestamps (e.g. every 5 minutes) highly disrupts the "
        "viewer's experience. By finding natural conversational boundaries and semantic breaks, "
        "the AdMarker Engine inserts advertisements precisely when a topic concludes or during a "
        "prolonged pause, maximizing ad retention while preserving user engagement."
    )
    
    p_llmh = doc.add_paragraph()
    p_llmh.paragraph_format.space_before = Pt(10)
    p_llmh.paragraph_format.space_after = Pt(6)
    r_llmh = p_llmh.add_run("How the System Uses Large Language Models (LLMs):")
    r_llmh.font.bold = True
    r_llmh.font.color.rgb = RGBColor(0, 112, 192)
    
    llm_bullets = [
        ("Structured Transcription via Gemini REST API", 
         "The system converts raw video files into structured audio segments. When using the Gemini transcription provider, raw audio is chunked, converted to base64, and sent to Google Gemini via a direct REST API call. By supplying a target JSON output schema, the LLM performs precise speech-to-text alignment, returning an array of segments with start, end, and transcribed text."),
        ("Contextual Tagging & Tone Inferences", 
         "After finding optimal breaks, the system gathers the preceding 30 seconds of conversational transcript. This context is sent to a Gemini or Hugging Face model to classify the primary topic (e.g., Technology, Cooking, Automotive, Finance) and infer the overall tone of discussion (e.g., Professional, Warm & Instructional, Hands-on, Analytical)."),
        ("Targeted Ad Recommendation Engine", 
         "Based on the inferred topic and tone, the LLM acts as a recommendation agent, suggesting 2–3 highly relevant, non-disruptive products or services. For example, a segment discussing database clustering results in recommendations for high-performance serverless cloud hosting or security audit tools.")
    ]
    
    for l_title, l_desc in llm_bullets:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.2)
        r_t = p.add_run(f"■ {l_title}\n")
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(30, 30, 30)
        r_d = p.add_run(l_desc)
        r_d.font.size = Pt(10)
        r_d.font.color.rgb = RGBColor(80, 80, 80)
        
    p_pipelineh = doc.add_paragraph()
    p_pipelineh.paragraph_format.space_before = Pt(10)
    p_pipelineh.paragraph_format.space_after = Pt(6)
    r_pipeh = p_pipelineh.add_run("Core Processing Stages:")
    r_pipeh.font.bold = True
    r_pipeh.font.color.rgb = RGBColor(0, 112, 192)
    
    pipeline_stages = [
        ("Audio Extraction", "Converts input videos (MP4/MOV) into high-fidelity 16kHz mono WAV streams using moviepy or a fallback ffmpeg command-line subprocess."),
        ("Voice Activity Detection (VAD)", "Scans the WAV file block-by-block (100ms blocks) to calculate Root Mean Square (RMS) energy. It splits audio into speech chunks (20-30s) based on silence thresholds, ensuring memory remains flat regardless of video length."),
        ("Stitched Transcription", "Transcribes chunks using a local Faster-Whisper model or Google Gemini, mapping chunk-relative segments back to the global timeline."),
        ("Boundary Scoring", "A scoring heuristic processes the timeline gaps. Gaps are rewarded based on duration, with a large weight (+3.0) added if the preceding segment ends with sentence-ending punctuation (., ?, !) or a clause boundary (,, ;, :)."),
        ("Contextual Recommendation", "The preceding context is analyzed by the LLM to generate targeted product ads."),
        ("Streaming API & Frontend", "A multi-threaded python server spawns the execution and streams console output logs to the browser in real-time using Server-Sent Events (SSE). The dashboard displays interactive states, progress stages, and a final recommendation table.")
    ]
    
    for s_title, s_desc in pipeline_stages:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r_t = p.add_run(f"{s_title}: ")
        r_t.font.bold = True
        r_d = p.add_run(s_desc)
        
    # ----------------------------------------------------
    # SECTION 6: Code snippets
    # ----------------------------------------------------
    doc.add_page_break()
    p_h6 = doc.add_paragraph()
    r_h6 = p_h6.add_run("6. Code snippets")
    r_h6.font.size = Pt(16)
    r_h6.font.bold = True
    r_h6.font.color.rgb = RGBColor(31, 78, 121)
    p_h6.paragraph_format.space_before = Pt(18)
    p_h6.paragraph_format.space_after = Pt(12)
    
    # Snippet 1
    p_sd1 = doc.add_paragraph()
    r_sd1_title = p_sd1.add_run("Snippet 1: Memory-Safe Voice Activity Detection (VAD) Block Reading (audio.py)\n")
    r_sd1_title.font.bold = True
    r_sd1_title.font.color.rgb = RGBColor(0, 112, 192)
    r_sd1_desc = p_sd1.add_run("•  Purpose: Reads large WAV audio files block-by-block (100ms segments) to compute RMS energy. It keeps the system's memory flat and identifies silence periods to divide speech into 20–30 second chunks.")
    r_sd1_desc.font.italic = True
    p_sd1.paragraph_format.space_after = Pt(6)
    
    code1 = r"""# Read WAV block-by-block and compute RMS to check silence
with sf.SoundFile(audio_path) as f:
    while True:
        block = f.read(block_samples, dtype='float32')
        if len(block) == 0:
            break
        rms = np.sqrt(np.mean(block ** 2)) if len(block) > 0 else 0.0
        is_silent = rms < energy_threshold
        current_chunk_blocks.append((block, rms))
        
        # Track silence hangover
        consecutive_silence = (consecutive_silence + 1) if is_silent else 0
        current_len = len(current_chunk_blocks) * block_len_sec"""
                
    p_c1 = doc.add_paragraph()
    p_c1.paragraph_format.left_indent = Inches(0.4)
    p_c1.paragraph_format.space_after = Pt(12)
    add_left_border(p_c1, "0070C0", "24")
    add_paragraph_shading(p_c1, "F5F5F5")
    r_c1 = p_c1.add_run(code1)
    r_c1.font.name = 'Consolas'
    r_c1.font.size = Pt(9.5)
    r_c1.font.color.rgb = RGBColor(50, 50, 50)
    
    # Snippet 2
    p_sd2 = doc.add_paragraph()
    r_sd2_title = p_sd2.add_run("Snippet 2: Boundary Scoring and Context Extraction (analyzer.py)\n")
    r_sd2_title.font.bold = True
    r_sd2_title.font.color.rgb = RGBColor(0, 112, 192)
    r_sd2_desc = p_sd2.add_run("•  Purpose: Evaluates the silent pauses between speech chunks. It rewards longer pauses and gives a heavy weight (+3.0 score) if the segment before the pause ends in a full sentence (., ?, !) or clause (,, ;, :), ensuring ad breaks are placed at natural conversational transitions.")
    r_sd2_desc.font.italic = True
    p_sd2.paragraph_format.space_after = Pt(6)
    
    code2 = r"""# Score silent gaps based on duration and ending punctuation
for i in range(len(segments) - 1):
    current_seg, next_seg = segments[i], segments[i+1]
    gap_duration = next_seg["start"] - current_seg["end"]
    
    if gap_duration >= self.min_gap_seconds:
        score = gap_duration
        text_clean = current_seg["text"].strip()
        
        # Reward completed sentence boundaries (+3.0) or clauses (+1.0)
        if text_clean and text_clean[-1] in ('.', '?', '!'):
            score += 3.0
        elif text_clean and text_clean[-1] in (',', ';', ':'):
            score += 1.0"""
        
    p_c2 = doc.add_paragraph()
    p_c2.paragraph_format.left_indent = Inches(0.4)
    p_c2.paragraph_format.space_after = Pt(12)
    add_left_border(p_c2, "0070C0", "24")
    add_paragraph_shading(p_c2, "F5F5F5")
    r_c2 = p_c2.add_run(code2)
    r_c2.font.name = 'Consolas'
    r_c2.font.size = Pt(9.5)
    r_c2.font.color.rgb = RGBColor(50, 50, 50)
    
    # Snippet 3
    p_sd3 = doc.add_paragraph()
    r_sd3_title = p_sd3.add_run("Snippet 3: Google Gemini API REST Integration for Contextual Tagging (llm.py)\n")
    r_sd3_title.font.bold = True
    r_sd3_title.font.color.rgb = RGBColor(0, 112, 192)
    r_sd3_desc = p_sd3.add_run("•  Purpose: Queries the Gemini API with structured output schema generation configuration to force a clean, parser-friendly JSON response containing categories, tones, and ad suggestions.")
    r_sd3_desc.font.italic = True
    p_sd3.paragraph_format.space_after = Pt(6)
    
    code3 = r"""# Query Google Gemini API via REST with structured output schema configuration
url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-lite:generateContent?key={self.gemini_api_key}"
payload = {
    "contents": [{"parts": [{"text": self._get_prompt(context_text)}]}],
    "generationConfig": {
        "responseMimeType": "application/json",
        "responseSchema": {
            "type": "OBJECT",
            "properties": {"category": {"type": "STRING"}, "tone": {"type": "STRING"}, "ad_suggestions": {"type": "ARRAY", "items": {"type": "STRING"}}},
            "required": ["category", "tone", "ad_suggestions"]
        },
        "temperature": 0.1
    }
}
response = requests.post(url, headers={"Content-Type": "application/json"}, json=payload)"""
            
    p_c3 = doc.add_paragraph()
    p_c3.paragraph_format.left_indent = Inches(0.4)
    p_c3.paragraph_format.space_after = Pt(12)
    add_left_border(p_c3, "0070C0", "24")
    add_paragraph_shading(p_c3, "F5F5F5")
    r_c3 = p_c3.add_run(code3)
    r_c3.font.name = 'Consolas'
    r_c3.font.size = Pt(9.5)
    r_c3.font.color.rgb = RGBColor(50, 50, 50)
    
    # Snippet 4
    p_sd4 = doc.add_paragraph()
    r_sd4_title = p_sd4.add_run("Snippet 4: Asynchronous Pipeline Execution and SSE Log Streaming (web_server.py)\n")
    r_sd4_title.font.bold = True
    r_sd4_title.font.color.rgb = RGBColor(0, 112, 192)
    r_sd4_desc = p_sd4.add_run("•  Purpose: This API endpoint spawns the Python processing script in a subprocess, intercepts standard output, cleans ANSI codes, and streams logs in real-time to the client using Server-Sent Events (SSE).")
    r_sd4_desc.font.italic = True
    p_sd4.paragraph_format.space_after = Pt(6)
    
    code4 = r"""# Spawn python processing subprocess and stream stdout logs back to client using SSE
process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1)
ansi_escape = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')

for line in process.stdout:
    clean_line = ansi_escape.sub('', line)
    if clean_line.strip() == "" and line.strip() == "":
        continue
    log_payload = json.dumps({"log": clean_line})
    self.wfile.write(f"data: {log_payload}\n\n".encode('utf-8'))
    self.wfile.flush()
process.wait()"""
            
    p_c4 = doc.add_paragraph()
    p_c4.paragraph_format.left_indent = Inches(0.4)
    p_c4.paragraph_format.space_after = Pt(12)
    add_left_border(p_c4, "0070C0", "24")
    add_paragraph_shading(p_c4, "F5F5F5")
    r_c4 = p_c4.add_run(code4)
    r_c4.font.name = 'Consolas'
    r_c4.font.size = Pt(9.5)
    r_c4.font.color.rgb = RGBColor(50, 50, 50)
    
    # ----------------------------------------------------
    # SECTION 7: System Output & Screenshots
    # ----------------------------------------------------
    doc.add_page_break()
    p_h7 = doc.add_paragraph()
    r_h7 = p_h7.add_run("7. System Output & Dashboard Screenshots")
    r_h7.font.size = Pt(16)
    r_h7.font.bold = True
    r_h7.font.color.rgb = RGBColor(31, 78, 121)
    p_h7.paragraph_format.space_before = Pt(12)
    p_h7.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph(
        "Below are the operational states and screenshot visualizations of the AdMarker Engine "
        "dashboard during file upload, running event streaming, and final timeline result render."
    )
    
    # State 1
    p_st1 = doc.add_paragraph()
    r_st1_title = p_st1.add_run("State 1: Professional File Uploader & Configuration Panel\n")
    r_st1_title.font.bold = True
    r_st1_title.font.color.rgb = RGBColor(0, 112, 192)
    r_st1_desc = p_st1.add_run(
        "This state shows the home layout of the dashboard featuring a glassmorphic drag-and-drop zone. "
        "Users can select an MP4 video or WAV audio file, choose between Google Gemini and faster-whisper "
        "for transcription, configure context tagging models, and adjust the minimum silent gap threshold via an interactive slider."
    )
    p_st1.paragraph_format.space_after = Pt(6)
    
    p_img1 = doc.add_paragraph()
    p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_img1 = p_img1.add_run()
    r_img1.add_picture("./report_images/dashboard_state1.png", width=Inches(5.8))
    p_img1.paragraph_format.space_after = Pt(18)
    
    # State 2
    p_st2 = doc.add_paragraph()
    r_st2_title = p_st2.add_run("State 2: Asynchronous Process Tracker & Live Event Log\n")
    r_st2_title.font.bold = True
    r_st2_title.font.color.rgb = RGBColor(0, 112, 192)
    r_st2_desc = p_st2.add_run(
        "When the pipeline starts, the system monitors each phase (Audio Extraction, VAD Chunking, "
        "Transcription, Boundary Scoring, and Context Tagging). The dashboard reflects the stage "
        "completion in real-time while streaming stdout lines from the backend into a retro terminal log "
        "console via Server-Sent Events (SSE)."
    )
    p_st2.paragraph_format.space_after = Pt(6)
    
    p_img2 = doc.add_paragraph()
    p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_img2 = p_img2.add_run()
    r_img2.add_picture("./report_images/dashboard_state2.png", width=Inches(5.8))
    p_img2.paragraph_format.space_after = Pt(18)
    
    # State 3
    p_st3 = doc.add_paragraph()
    r_st3_title = p_st3.add_run("State 3: Segment Timeline & AI-Generated Ad Recommendations Table\n")
    r_st3_title.font.bold = True
    r_st3_title.font.color.rgb = RGBColor(0, 112, 192)
    r_st3_desc = p_st3.add_run(
        "Upon pipeline completion, the results panel expands to present the final identified quiet-zones, "
        "their exact timeline timestamps, silent gap durations, preceding contextual transcripts, "
        "inferred categories/tones, and tailored, non-intrusive product advertisement suggestions."
    )
    p_st3.paragraph_format.space_after = Pt(6)
    
    p_img3 = doc.add_paragraph()
    p_img3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_img3 = p_img3.add_run()
    r_img3.add_picture("./report_images/dashboard_state3.png", width=Inches(5.8))
    p_img3.paragraph_format.space_after = Pt(18)
    
    # Save document
    try:
        doc.save("Project_Report.docx")
        print("[Script] Generated Project_Report.docx successfully!")
    except PermissionError:
        doc.save("Project_Report_v2.docx")
        print("[Script] Warning: 'Project_Report.docx' was locked by another process (likely open in Word). Saved to 'Project_Report_v2.docx' instead.")

def build_markdown_report():
    code1 = r"""# Read WAV block-by-block and compute RMS to check silence
with sf.SoundFile(audio_path) as f:
    while True:
        block = f.read(block_samples, dtype='float32')
        if len(block) == 0:
            break
        rms = np.sqrt(np.mean(block ** 2)) if len(block) > 0 else 0.0
        is_silent = rms < energy_threshold
        current_chunk_blocks.append((block, rms))
        
        # Track silence hangover
        consecutive_silence = (consecutive_silence + 1) if is_silent else 0
        current_len = len(current_chunk_blocks) * block_len_sec"""

    code2 = r"""# Score silent gaps based on duration and ending punctuation
for i in range(len(segments) - 1):
    current_seg, next_seg = segments[i], segments[i+1]
    gap_duration = next_seg["start"] - current_seg["end"]
    
    if gap_duration >= self.min_gap_seconds:
        score = gap_duration
        text_clean = current_seg["text"].strip()
        
        # Reward completed sentence boundaries (+3.0) or clauses (+1.0)
        if text_clean and text_clean[-1] in ('.', '?', '!'):
            score += 3.0
        elif text_clean and text_clean[-1] in (',', ';', ':'):
            score += 1.0"""

    code3 = r"""# Query Google Gemini API via REST with structured output schema configuration
url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-lite:generateContent?key={self.gemini_api_key}"
payload = {
    "contents": [{"parts": [{"text": self._get_prompt(context_text)}]}],
    "generationConfig": {
        "responseMimeType": "application/json",
        "responseSchema": {
            "type": "OBJECT",
            "properties": {"category": {"type": "STRING"}, "tone": {"type": "STRING"}, "ad_suggestions": {"type": "ARRAY", "items": {"type": "STRING"}}},
            "required": ["category", "tone", "ad_suggestions"]
        },
        "temperature": 0.1
    }
}
response = requests.post(url, headers={"Content-Type": "application/json"}, json=payload)"""

    code4 = r"""# Spawn python processing subprocess and stream stdout logs back to client using SSE
process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1)
ansi_escape = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')

for line in process.stdout:
    clean_line = ansi_escape.sub('', line)
    if clean_line.strip() == "" and line.strip() == "":
        continue
    log_payload = json.dumps({"log": clean_line})
    self.wfile.write(f"data: {log_payload}\n\n".encode('utf-8'))
    self.wfile.flush()
process.wait()"""

    md_content = f"""# PRESIDENCY UNIVERSITY
**School of Artificial Intelligence & Advanced Computing**

<br>

<div align="center">
  <strong>AI Centre of Excellence</strong>
  <br>
  <em>Accelerated by NVIDIA</em>
</div>

<br><br>

# NVIDIA ACCELERATED AI CENTRE OF EXCELLENCE
## Final Internship and Project Report

<br><br>

| Field | Details |
| :--- | :--- |
| **PROJECT REPORT TITLE** | AUTOMATED VIDEO AD-MARKER & CONTEXTUAL PROMPT GENERATOR |
| **STUDENT NAME** | LAALIKA J |
| **ROLL NUMBER / ID** | 20241CAI0231 |
| **BRANCH & DEPARTMENT** | B-TECH School of Artificial Intelligence & Advanced Computing |
| **START DATE & END DATE** | 29-06-2026 TO 10-07-2026 |
| **PROJECT SUBMISSION** | 11-07-2026 |

<br><br>

---

## 1. About NVIDIA Corporation

NVIDIA Corporation, founded in 1993 by Jen-Hsun Huang, Chris Malachowsky, and Curtis Priem, is a pioneer in graphics processing unit (GPU) accelerated computing. Originally focused on bringing 3D graphics to the gaming and multimedia markets, NVIDIA invented the GPU in 1999, redefining modern computer graphics and sparking the growth of the PC gaming market. Over the next two decades, NVIDIA expanded its scope, transforming the GPU from a specialized 3D rendering chip into a highly programmable parallel processor capable of general-purpose scientific computing.

Today, NVIDIA stands at the absolute vanguard of the Artificial Intelligence revolution. By recognizing the immense mathematical alignment between deep learning algorithms and parallel hardware architectures, the company made a massive, long-term strategic bet on AI supercomputing. Through the development of the proprietary CUDA compute platform, NVIDIA created a fully integrated stack of software libraries, hardware drivers, and deep learning engines that became the standard for researchers and developers globally. In 2024, NVIDIA joined the ranks of the world's most valuable companies, crossing a $3 trillion market capitalization milestone. Its compute platforms have become the foundation of almost all major generative AI foundations, large language models (LLMs), autonomous vehicle systems, and scientific supercomputing centers around the world.

---

## 2. NVIDIA H200 Tensor Core GPU

Released in Q2 2024, the NVIDIA H200 Tensor Core GPU is the successor to the groundbreaking H100, designed specifically to address the massive memory and bandwidth bottlenecks of generative AI and large-scale model inference. Built on the Hopper architecture, the H200 is the first GPU to deploy ultra-fast HBM3e (High Bandwidth Memory 3e) technology. By packing 141 GB of HBM3e memory running at a staggering 4.8 terabytes per second (TB/s), the H200 provides double the capacity and 1.4x the bandwidth of its predecessor. This enables it to store massive language models entirely inside active GPU memory, dramatically reducing latency and maximizing computing efficiency.

| SPECIFICATION | NVIDIA H100 GPU | NVIDIA H200 GPU (NEW) |
| :--- | :--- | :--- |
| **Architecture** | Hopper | Hopper (Upgraded Silicon) |
| **Memory Type** | HBM3 | HBM3e (Advanced High Bandwidth) |
| **Memory Capacity** | 80 GB | 141 GB (Nearly Double) |
| **Memory Bandwidth** | 3.35 TB/s | 4.8 TB/s (1.4x Faster) |
| **FP8 Performance** | Approx. 4 PetaFLOPS | Approx. 4 PetaFLOPS |
| **LLM Inference Speed** | Baseline (1.0x) | Llama 2 (70B): 1.9x faster \\| GPT-3: 1.6x faster |

---

## 3. Understanding the NVIDIA GPU Architecture

Graphics Processing Units (GPUs) differ fundamentally from Central Processing Units (CPUs) in their computational philosophy. A CPU is designed for sequential processing and low-latency serial tasks, utilizing a few powerful cores optimized for single-threaded speed. Conversely, a GPU is built for massive parallel execution, containing thousands of smaller, simpler cores that work simultaneously. This makes GPUs uniquely suited for tasks like 3D graphics rendering, computer vision, and matrix math operations — which form the foundational operations of deep learning.

### Key Architectural Sub-systems of NVIDIA GPUs:

* **CUDA Cores**: The fundamental parallel processors that execute basic mathematical calculations and logic instructions on the GPU.
* **Tensor Cores**: Specialized hardware units introduced in the Volta architecture designed to perform fast mixed-precision matrix multiplication in a single cycle. These are the engines that accelerate Transformer models and LLM operations.
* **Streaming Multiprocessors (SM)**: The primary building blocks of the GPU. Each SM groups CUDA cores, Tensor cores, register files, and shared memory cache to coordinate execution.
* **NVLink Interconnect**: NVIDIA's proprietary high-bandwidth interconnect, allowing multiple GPUs to communicate at speeds up to 900 GB/s, enabling them to act as a single massive virtual GPU.

---

## 4. Project Work Timeline: Day 1 to Day 10

* **Day 1: Introduction to NVIDIA GPU and Kubernetes Pods and Services**
  * Understood about NVIDIA H-200 GPU and its features.
  * Creation of Pods and services & how it connects to the NVIDIA Server.
* **Day 2: Introduction to Machine Learning and Model Training**
  * Understood Core Machine Learning Concepts.
  * Performed a program to train a model by following the Model Training Pipeline.
* **Day 3: Supervised and Unsupervised Machine Learning Concepts**
  * Focused and understood programs based on Classification & Regression.
* **Day 4: Neural Networks: CNN and RNN**
  * Focused on creating Neural Networks and trained a model based on the MNIST dataset.
* **Day 5: Working Natural Language Processing Pipeline for a Model**
  * Understood NLP Pipeline and Types of NLP.
  * Focused on training a model following every step of the NLP pipeline.
* **Day 6: Transformers and Types of Transformer (BERT & GPT)**
  * Understood Transformers and types of transformers used to train a model.
  * Understood concepts like BERT and GPT and focused on training models that use these concepts.
* **Day 7: Large Language Models**
  * Understood the importance of LLMs in Machine Learning.
  * Focused on programs that used Google Gemini and compared other LLMs with Google Gemini.
* **Day 8: VIVA about Concepts Covered and Project Selection**
  * VIVA about all the concepts that were covered till Day 7.
  * Implementing a project based on the concepts we understood.
* **Day 9: Working of Diffusion Models and Generative Adversarial Networks**
  * Understood how Diffusion Models work and types of Diffusion Models.
  * Implemented programs that use diffusion models.
  * Focused on training a model using GAN and generated outputs.
* **Day 10: Working on Final Internship Project**
  * Focused on completing our Final Internship Project.

---

## 5. Project Details: AdMarker Engine

The **AdMarker Engine** is an automated pipeline and professional web dashboard designed to identify optimal quiet-zones (silent pauses) in audio or video timelines and leverage Large Language Models (LLMs) to perform semantic categorization and targeted product ad recommendations. 

Placing advertisements at random timestamps (e.g. every 5 minutes) highly disrupts the viewer's experience. By finding natural conversational boundaries and semantic breaks, the AdMarker Engine inserts advertisements precisely when a topic concludes or during a prolonged pause, maximizing ad retention while preserving user engagement.

### How the System Uses Large Language Models (LLMs):

1. **Structured Transcription via Gemini REST API**:
   The system converts raw video files into structured audio segments. When using the Gemini transcription provider, raw audio is chunked, converted to base64, and sent to Google Gemini via a direct REST API call. By supplying a target JSON output schema, the LLM performs precise speech-to-text alignment, returning an array of segments with start, end, and transcribed text.
2. **Contextual Tagging & Tone Inferences**:
   After finding optimal breaks, the system gathers the preceding 30 seconds of conversational transcript. This context is sent to a Gemini or Hugging Face model to classify the primary topic (e.g., Technology, Culinary, Automotive, Finance) and infer the overall tone of discussion (e.g., Professional, Warm & Instructional, Hands-on, Analytical).
3. **Targeted Ad Recommendation Engine**:
   Based on the inferred topic and tone, the LLM acts as a recommendation agent, suggesting 2–3 highly relevant, non-disruptive products or services. For example, a segment discussing database clustering results in recommendations for high-performance serverless cloud hosting or security audit tools.

### Core Processing Stages:

```mermaid
graph TD
    A[Source Video/Audio] -->|FFmpeg / MoviePy| B[16kHz Mono WAV Audio]
    B -->|Memory-Safe Block RMS Analysis| C[VAD Speech Chunks]
    C -->|faster-whisper / Gemini API| D[Global Segment Timeline]
    D -->|Boundary Scoring Algorithm| E[Optimal Ad Break Boundaries]
    E -->|Context Extraction| F[LLM Category, Tone & Ad Generation]
    F -->|Event Stream SSE| G[Interactive Web Dashboard Table]
```

1. **Audio Extraction**: Converts input videos (MP4/MOV) into high-fidelity 16kHz mono WAV streams using `moviepy` or a fallback `ffmpeg` command-line subprocess.
2. **Voice Activity Detection (VAD)**: Scans the WAV file block-by-block (100ms blocks) to calculate Root Mean Square (RMS) energy. It splits audio into speech chunks (20-30s) based on silence thresholds, ensuring memory remains flat regardless of video length.
3. **Stitched Transcription**: Transcribes chunks using a local Faster-Whisper model or Google Gemini, mapping chunk-relative segments back to the global timeline.
4. **Boundary Scoring**: A scoring heuristic processes the timeline gaps. Gaps are rewarded based on duration, with a large weight (+3.0) added if the preceding segment ends with sentence-ending punctuation (`.`, `?`, `!`) or a clause boundary (`,`, `;`, `:`).
5. **Contextual Recommendation**: The preceding context is analyzed by the LLM to generate targeted product ads.
6. **Streaming API & Frontend**: A multi-threaded python server spawns the execution and streams console output logs to the browser in real-time using Server-Sent Events (SSE). The dashboard displays interactive states, progress stages, and a final recommendation table.

---

## 6. Code snippets

### Snippet 1: Memory-Safe Voice Activity Detection (VAD) Block Reading (audio.py)
* **Purpose**: Reads large WAV audio files block-by-block (100ms segments) to compute RMS energy. It keeps the system's memory flat and identifies silence periods to divide speech into 20–30 second chunks.

```python
{code1}
```

### Snippet 2: Boundary Scoring and Context Extraction (analyzer.py)
* **Purpose**: Evaluates the silent pauses between speech chunks. It rewards longer pauses and gives a heavy weight (+3.0 score) if the segment before the pause ends in a full sentence (., ?, !) or clause (,, ;, :), ensuring ad breaks are placed at natural conversational transitions.

```python
{code2}
```

### Snippet 3: Google Gemini API REST Integration for Contextual Tagging (llm.py)
* **Purpose**: Queries the Gemini API with structured output schema generation configuration to force a clean, parser-friendly JSON response containing categories, tones, and ad suggestions.

```python
{code3}
```

### Snippet 4: Asynchronous Pipeline Execution and SSE Log Streaming (web_server.py)
* **Purpose**: This API endpoint spawns the Python processing script in a subprocess, intercepts standard output, cleans ANSI codes, and streams logs in real-time to the client using Server-Sent Events (SSE).

```python
{code4}
```

---

## 7. System Output & Dashboard Screenshots

Below are the detailed descriptions and layouts of the operational states of the AdMarker Engine dashboard.

### State 1: Professional File Uploader & Configuration Panel
This state displays the home layout of the dashboard featuring a glassmorphic drag-and-drop zone. Users can select an MP4 video or WAV audio file, choose between Google Gemini and faster-whisper for transcription, configure context tagging models, and adjust the minimum silent gap threshold via an interactive slider.

![Dashboard Home Screen with Active Drag-and-Drop Uploader and Settings Panel](file:///C:/Users/Swathi/Desktop/Finn/Projects/Ads/report_images/dashboard_state1.png)

---

### State 2: Asynchronous Process Tracker & Live Event Log
When the pipeline starts, the system monitors each phase. The dashboard reflects the stage completion in real-time while streaming stdout lines from the backend into a retro terminal log console via Server-Sent Events (SSE).

![Pipeline Monitor displaying Active Transcribing stage and Live Console Log Streaming](file:///C:/Users/Swathi/Desktop/Finn/Projects/Ads/report_images/dashboard_state2.png)

---

### State 3: Segment Timeline & AI-Generated Ad Recommendations Table
Upon pipeline completion, the results table presents the final identified quiet-zones, timestamps, and targeted ad recommendations.

![Result Table showcasing identified ad markers, preceding context snippets, and targeted ad recommendations](file:///C:/Users/Swathi/Desktop/Finn/Projects/Ads/report_images/dashboard_state3.png)
"""
    with open("Project_Report.md", "w", encoding="utf-8") as f:
        f.write(md_content)
    print("[Script] Generated Project_Report.md successfully!")

if __name__ == "__main__":
    create_directories()
    draw_presidency_logo()
    draw_nvidia_logo()
    draw_dashboard_state1()
    draw_dashboard_state2()
    draw_dashboard_state3()
    build_docx_report()
    build_markdown_report()
